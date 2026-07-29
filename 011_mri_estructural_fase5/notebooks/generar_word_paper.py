"""
Genera el bosquejo de la publicación en formato Word editable (.docx).

Inserta las figuras ya producidas en el lugar que les corresponde y deja
**marcadores visibles** donde faltan figuras por generar (las anatómicas con
FreeSurfer, ver `docs/PROMPTS_PARA_RETOMAR.md` prompt B).

El .docx NO se versiona (el .gitignore excluye *.docx por política de PII, y
además es un binario regenerable). Se escribe en `docs/` y se copia a la carpeta
de documentación del proyecto en OneDrive.

Ejecutar:  ~/FS_FONDECYT/.venv/bin/python notebooks/generar_word_paper.py
"""

# %% ── setup ────────────────────────────────────────────────────────────────
import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import config as cfg

FIGS = cfg.FIGS
SALIDA = cfg.DOCS / "PAPER_BOSQUEJO.docx"

TINTA = RGBColor(0x0B, 0x0B, 0x0B)
TINTA2 = RGBColor(0x52, 0x51, 0x4E)
ACENTO = RGBColor(0xEB, 0x68, 0x34)

doc = Document()

# Estilo base
est = doc.styles["Normal"]
est.font.name = "Calibri"
est.font.size = Pt(11)
est.paragraph_format.space_after = Pt(8)
est.paragraph_format.line_spacing = 1.15


# %% ── utilidades ───────────────────────────────────────────────────────────
def sombrear(celda, hex_color):
    tc = celda._tc.get_or_add_tcPr()
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:fill"), hex_color)
    tc.append(sombra)


def titulo(texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for r in h.runs:
        r.font.color.rgb = TINTA
    return h


def parrafo(texto, cursiva=False, size=11, color=TINTA):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.italic = cursiva
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def nota(texto):
    """Bloque destacado, para cautelas y decisiones."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    sombrear(c, "FDF3EE")
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(texto)
    r.font.size = Pt(10)
    r.font.color.rgb = TINTA2
    doc.add_paragraph()
    return t


def figura(ruta_png, numero, titulo_fig, pie, ancho_cm=15.5):
    """Inserta una figura existente con su leyenda numerada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(ruta_png).exists():
        p.add_run().add_picture(str(ruta_png), width=Cm(ancho_cm))
    else:
        r = p.add_run(f"[falta el archivo: {ruta_png}]")
        r.font.color.rgb = ACENTO
    leyenda = doc.add_paragraph()
    leyenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = leyenda.add_run(f"Figura {numero}. {titulo_fig} ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r2 = leyenda.add_run(pie)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = TINTA2
    doc.add_paragraph()


def hueco_figura(numero, titulo_fig, que_debe_mostrar, como_generarla):
    """Marcador visible para una figura que aún no existe."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    sombrear(c, "FFF6E5")
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"\n▢  ESPACIO RESERVADO — FIGURA {numero}\n")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACENTO
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(titulo_fig + "\n")
    r2.bold = True
    r2.font.size = Pt(10.5)
    p3 = c.add_paragraph()
    r3 = p3.add_run("Qué debe mostrar: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    r3b = p3.add_run(que_debe_mostrar)
    r3b.font.size = Pt(9.5)
    p4 = c.add_paragraph()
    r4 = p4.add_run("Cómo generarla: ")
    r4.bold = True
    r4.font.size = Pt(9.5)
    r4b = p4.add_run(como_generarla + "\n")
    r4b.font.size = Pt(9.5)
    r4b.font.color.rgb = TINTA2
    doc.add_paragraph()


def tabla(encabezados, filas, anchos=None, resaltar_filas=()):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(encabezados):
        celda = t.rows[0].cells[i]
        celda.text = ""
        r = celda.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(9)
    for j, fila in enumerate(filas):
        cells = t.add_row().cells
        for i, v in enumerate(fila):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            if j in resaltar_filas:
                r.bold = True
    doc.add_paragraph()
    return t


# %% ── PORTADA ──────────────────────────────────────────────────────────────
h = doc.add_heading(
    "Girificación cortical reducida como marcador de rasgo y adelgazamiento "
    "cortical como marcador de estado en el Mareo Postural Perceptual Persistente",
    level=0)
for r in h.runs:
    r.font.color.rgb = TINTA

parrafo("Un estudio de morfometría estructural con FreeSurfer", cursiva=True, size=12,
        color=TINTA2)
parrafo("FONDECYT de Iniciación 11200469 · Fase 011 · Bosquejo editable, 2026-07-28",
        size=10, color=TINTA2)

nota("CÓMO USAR ESTE DOCUMENTO. Es un bosquejo editable, no un manuscrito final. "
     "Las seis figuras están generadas y son definitivas; las de superficie se "
     "renderizaron con FreeSurfer. "
     "La sección de Métodos aquí es una versión abreviada; la versión exhaustiva, con la "
     "justificación de cada decisión analítica, está en docs/PAPER_BORRADOR.md §2. "
     "Todos los números provienen de results/ y son trazables.")

doc.add_page_break()

# %% ── RESUMEN ──────────────────────────────────────────────────────────────
titulo("Resumen", 1)

parrafo("Antecedentes. El Mareo Postural Perceptual Persistente (MPPP/PPPD) es un trastorno "
        "vestibular funcional crónico que aparece tras un evento vestibular agudo en solo una "
        "fracción de quienes lo sufren. No se sabe qué distingue a quien cronifica de quien se "
        "recupera, ni si esa diferencia es previa al evento o consecuencia de él.")

parrafo("Métodos. Morfometría de superficie con FreeSurfer en 46 sujetos (17 MPPP, 19 "
        "pacientes vestibulares sin cronificación, 10 voluntarios sanos). Cuatro medidas "
        "—índice de girificación local (LGI), grosor cortical, área y volumen— sobre 19 "
        "regiones de interés congeladas a priori, más barrido whole-brain por región y "
        "vertex-wise. El diseño principal fue un contraste dirigido MPPP vs vestibular. Se "
        "evaluó la asociación con navegación espacial alocéntrica (CSE, Entropy-Ratio) y con "
        "severidad sintomática (Niigata, DHI).")

parrafo("Resultados. Se observó una doble disociación. El LGI diferenció grupos (índice de red "
        "d = −0,94; seis regiones con d entre −0,85 y −1,04) pero no se asoció con ninguna "
        "medida conductual ni clínica (0 de 260 correlaciones sobrevivieron a la corrección). "
        "El grosor cortical hizo lo contrario: ninguna diferencia entre grupos en 136 pruebas "
        "dirigidas, pero 32 de 260 correlaciones con la severidad sobrevivieron (ρ hasta "
        "−0,70). El análisis vertex-wise replicó ambos patrones de forma independiente. El "
        "efecto del LGI se reprodujo en las tres parcelaciones corticales (r = 0,95–0,997) y "
        "ningún hallazgo dependió de un solo sujeto.")

parrafo("Conclusión. La girificación cortical, que se establece en el desarrollo temprano y es "
        "estable en la adultez, se comporta como marcador de rasgo asociado a la "
        "cronificación; el grosor cortical, plástico, como marcador de estado asociado a la "
        "gravedad actual. La limitación central es que el contraste informativo se da frente a "
        "pacientes vestibulares y no frente a sanos, lo que impide determinar la dirección "
        "del efecto.")

doc.add_page_break()

# %% ── INTRODUCCIÓN ─────────────────────────────────────────────────────────
titulo("1. Introducción", 1)

nota("SECCIÓN A COMPLETAR. El texto siguiente es el esqueleto argumental derivado de los "
     "datos. Debe enriquecerse con la revisión de literatura del proyecto "
     "(02_Revision_Literatura_Areas_Cerebrales_PPPD.md), especialmente: criterios Bárány, "
     "modelos de reponderación maladaptativa, y los antecedentes de morfometría en PPPD "
     "(incluido Nigro et al. sobre girificación).")

parrafo("El MPPP se caracteriza por mareo no vertiginoso, inestabilidad perceptual y "
        "sensibilidad a estímulos visuales complejos, persistiendo tres meses o más tras un "
        "evento vestibular desencadenante. Su fisiopatología se ha descrito como una "
        "reponderación maladaptativa entre las señales visual, vestibular y propioceptiva, "
        "con dependencia visual aumentada y control postural rígido.")

parrafo("Dos observaciones motivan este trabajo. Primera: no todo el que sufre un evento "
        "vestibular agudo cronifica, lo que sugiere factores de vulnerabilidad previos; pero "
        "el grupo de comparación adecuado para investigarlos no es el sujeto sano, sino el "
        "paciente vestibular que no cronificó. Segunda: la girificación cortical es un rasgo "
        "del desarrollo, establecido perinatalmente y estable en la adultez, de modo que una "
        "diferencia en esta medida difícilmente puede ser consecuencia de un cuadro de meses "
        "de evolución. El grosor cortical, en cambio, sí responde a procesos adquiridos.")

parrafo("Se preinscribió una lista de 19 regiones de interés de la red de navegación "
        "visuoespacial-vestibular, congelada antes de examinar ningún resultado, y se "
        "preespecificó el plan estadístico completo.")

figura(FIGS / "etapaD2_freesurfer" / "fig1_mapa_red_DCNN.png", 1,
       "La red de navegación visuoespacial-vestibular estudiada.",
       "Las regiones de interés congeladas a priori, sobre la superficie inflada de "
       "fsaverage. Naranja: prioridad alta (ínsula posterior, supramarginal, temporal "
       "superior, parahipocampal, entorrinal, precúneo, istmo del cíngulo). Azul: prioridad "
       "media. Las regiones subcorticales incluidas en la hipótesis (hipocampo, tálamo, "
       "amígdala, cerebelo) no tienen representación en superficie. Renderizado con "
       "FreeSurfer; el dibujo usa Desikan-Killiany, cuyos límites son casi idénticos a los "
       "de DKT empleado en el análisis (r = 0,997).")

doc.add_page_break()

# %% ── MÉTODOS ──────────────────────────────────────────────────────────────
titulo("2. Métodos", 1)

titulo("2.1 Participantes", 2)
parrafo("De una cohorte inicial de 53 sujetos se excluyeron seis por ausencia o corrupción de "
        "imagen cruda y uno por control de calidad (sub-extensión pial global bilateral "
        "detectada en inspección visual). Un sujeto adicional careció de LGI utilizable.")
tabla(["Grupo", "n", "Definición"],
      [["MPPP", "17", "Criterios Bárány de PPPD"],
       ["Vestibular", "19", "Patología vestibular periférica documentada, sin cronificación"],
       ["Voluntario Sano", "10", "Sin antecedente vestibular"],
       ["Total", "46", "N = 45 en las medidas de LGI"]],
      resaltar_filas=(3,))

titulo("2.2 El contraste dirigido: justificación previa al resultado", 2)
parrafo("El diseño de tres grupos está limitado por el brazo sano (n = 10). Se preespecificó "
        "un contraste dirigido MPPP vs Vestibular (n = 36) por dos razones independientes. "
        "Teórica: ambos grupos comparten historia de patología vestibular, y lo que los separa "
        "es la cronificación perceptual, que es la pregunta del estudio; comparar MPPP con "
        "sanos confunde dos efectos. Estadística: elimina la dependencia del brazo de n = 10 y "
        "produce dos grupos balanceados.")
nota("Este contraste no es una búsqueda posterior de significación. Se verificó que estima el "
     "MISMO efecto que el diseño de tres grupos, con mayor precisión: la correlación entre los "
     "tamaños de efecto de ambos diseños, sobre las 136 pruebas, es r = 0,99. Ambos diseños se "
     "reportan completos.")

titulo("2.3 Procesamiento de imagen", 2)
parrafo("Reconstrucción cortical completa con FreeSurfer 8.2.0 (recon-all -all -T2pial "
        "-qcache). El flag -T2pial refina la superficie pial con la imagen T2, reduciendo la "
        "sobre-extensión hacia duramadre. Control de calidad por inspección visual, más un "
        "control cuantitativo mediante SurfaceHoles (número de agujeros topológicos reparados, "
        "relacionado con el número de Euler), que no difiere entre grupos (p = 0,896): esto "
        "descarta que los hallazgos sean artefacto de calidad diferencial de reconstrucción.")

nota("CORRECCIÓN A DECLARAR EN METHODS. Se detectó que el script de extracción leía la "
     "desviación estándar intrarregional del LGI en lugar del LGI medio. Los valores se "
     "recalcularon desde los archivos .stats originales; el rango resultante (1,64–4,98; "
     "mediana 2,71) corresponde al rango fisiológico esperado.")

titulo("2.4 Modelo estadístico", 2)
parrafo("Para cada combinación de región, hemisferio y medida se ajustó: medida ~ Grupo + Edad "
        "+ Sexo + Nivel educacional, añadiendo eTIV en volumen y área. Las covariables no se "
        "eligieron por costumbre: cada candidata se contrastó entre grupos y ninguna difiere "
        "(edad p = 0,816; sexo 0,131; educación 0,749; eTIV 0,535; SurfaceHoles 0,896; "
        "lateralidad 0,854). Los grupos están bien emparejados, de modo que el ajuste aumenta "
        "precisión pero no corrige sesgo de confusión.")

parrafo("El valor p reportado es el de permutación (Freedman-Lane, 10.000 remuestreos), no el "
        "paramétrico: con n = 10 en un brazo la distribución F asintótica es poco fiable. El "
        "procedimiento permuta los residuos del modelo sin el efecto de grupo, destruyendo ese "
        "efecto pero conservando la estructura de las covariables. Se conservan en las tablas "
        "los valores paramétrico, robusto (HC3) y de Kruskal-Wallis para comprobar que las "
        "conclusiones no dependen del método.")

parrafo("Los tamaños de efecto se reportan siempre con intervalo de confianza del 95% "
        "bootstrap BCa estratificado por grupo. La corrección por comparaciones múltiples es "
        "FDR de Benjamini-Hochberg dentro de cada familia (una medida dentro de una etapa), "
        "nunca sobre el conjunto total: los tres atlas corticales son tres parcelaciones del "
        "mismo manto y sus columnas están fuertemente correlacionadas.")

titulo("2.5 Análisis vertex-wise", 2)
parrafo("El mismo modelo ajustado en cada uno de los ~164.000 vértices de fsaverage "
        "(mri_glmfit --doss), con corrección por clusters mediante simulación de Monte Carlo: "
        "umbral de formación p < 0,001, umbral corregido de cluster CWP < 0,05, y corrección "
        "adicional por los dos hemisferios.")

nota("PUNTO TÉCNICO RELEVANTE PARA REPLICACIÓN. Al aplicar el suavizado estándar de 10 mm al "
     "LGI, la simulación falló: el FWHM residual alcanzó 37 mm, fuera del rango de las tablas "
     "precomputadas de FreeSurfer. El LGI, al integrar un parche de superficie amplio por "
     "construcción, ya llega suavizado: sin añadir nada tiene un FWHM residual de 10,5, "
     "comparable al del grosor con fwhm10 (14,5). Se analizó por tanto sin suavizado "
     "adicional. El fallo es silencioso, de modo que conviene declararlo.")

doc.add_page_break()

# %% ── RESULTADOS ───────────────────────────────────────────────────────────
titulo("3. Resultados", 1)

titulo("3.1 Características de la muestra", 2)
parrafo("Los grupos son demográficamente indistinguibles y clínicamente muy distintos: MPPP "
        "presenta mayor severidad, mayor ansiedad rasgo, peor cribado cognitivo y peor "
        "desempeño en navegación alocéntrica.")
tabla(["Variable", "N", "Sano", "Vestibular", "MPPP", "p"],
      [["Edad (años)", "46", "43,0", "45,0", "48,0", "0,816"],
       ["Sexo (F/M)", "46", "5/5", "16/3", "13/4", "0,131"],
       ["eTIV (×10⁶ mm³)", "46", "1,67", "1,53", "1,57", "0,535"],
       ["SurfaceHoles", "46", "29,0", "26,0", "25,0", "0,896"],
       ["DHI", "35", "1,0", "34,0", "46,0", "0,015"],
       ["Niigata", "35", "1,0", "17,0", "30,0", "0,004"],
       ["STAI-Rasgo", "34", "20,0", "23,0", "27,0", "0,004"],
       ["MoCA", "45", "26,0", "27,0", "24,0", "0,019"],
       ["CSE", "46", "28,0", "38,0", "60,9", "0,016"],
       ["Entropy-Ratio", "46", "0,5", "0,5", "0,6", "0,006"]],
      resaltar_filas=(4, 5, 6, 7, 8, 9))
parrafo("Tabla 1. Mediana por grupo; contraste de Kruskal-Wallis, o χ² para sexo. "
        "Los rangos intercuartílicos completos están en results/"
        "etapa0_tabla1_resultados_descriptivos.csv.", size=9, color=TINTA2)

titulo("3.2 Diferencias entre grupos: solo la girificación", 2)
parrafo("De 182 pruebas en las regiones preinscritas, ninguna sobrevivió al FDR en el diseño "
        "de tres grupos. Tampoco lo hizo ninguna de las 2.268 pruebas del barrido whole-brain "
        "por tabla. El contraste dirigido sí produjo resultados, y todos son LGI.")

tabla(["Región", "Hemi", "Medida", "d", "IC 95%", "p FDR"],
      [["Índice de red DCNN", "bilat", "LGI", "−0,94", "−1,73 a −0,03", "0,040"],
       ["Ínsula posterior", "der", "LGI", "−1,04", "−1,73 a −0,16", "0,043"],
       ["Temporal superior", "der", "LGI", "−1,04", "−1,74 a −0,23", "0,043"],
       ["Temporal superior", "izq", "LGI", "−0,94", "−1,70 a −0,13", "0,047"],
       ["Ínsula posterior", "izq", "LGI", "−0,89", "−1,63 a −0,02", "0,047"],
       ["Giro supramarginal", "izq", "LGI", "−0,88", "−1,55 a −0,07", "0,047"],
       ["Parahipocampal", "izq", "LGI", "−0,85", "−1,66 a +0,13", "0,047"]],
      resaltar_filas=(0,))
parrafo("Tabla 2. Todo lo que sobrevive a la corrección en el contraste dirigido "
        "MPPP vs Vestibular. Ninguna prueba de grosor, área o volumen sobrevivió en ningún "
        "diseño.", size=9, color=TINTA2)

figura(FIGS / "sintesis" / "forest_LGI_todas_las_rois.png", 2,
       "Girificación en las 32 regiones a priori.",
       "Contraste dirigido MPPP vs Vestibular, n = 36. Cada línea es una región con su "
       "intervalo de confianza del 95% (bootstrap BCa). Treinta y una de treinta y dos "
       "regiones caen del lado negativo: la girificación es menor en MPPP con notable "
       "consistencia. En negrita, las que sobreviven al FDR de su familia.",
       ancho_cm=13.0)

parrafo("El barrido whole-brain, aunque no declara ninguna región, muestra dónde se concentra "
        "la señal: el LGI acumula 73 pruebas con p < 0,05 cuando se esperarían 13,9 por azar "
        "(5,3×), mientras el grosor acumula 3 cuando se esperarían 13,9 (0,2×, menos que el "
        "azar). En z-scores promediados sobre las regiones de prioridad alta, el gradiente es "
        "MPPP −0,28 · Sano +0,04 · Vestibular +0,24.")

figura(FIGS / "etapaD2_freesurfer" / "fig2_clusters_LGI.png", 3,
       "Clusters de girificación corregidos por comparaciones múltiples.",
       "Análisis vertex-wise sobre ~164.000 vértices por hemisferio, corregido por clusters "
       "(Monte Carlo, umbral p < 0,001, CWP < 0,05, corregido por dos hemisferios). Azul = "
       "menor en MPPP. Se muestra únicamente el mapa enmascarado por los clusters que "
       "sobreviven. Un procedimiento que desconoce la lista preinscrita de regiones converge "
       "en la misma medida, dirección y contraste.")

titulo("3.3 Asociación con conducta y clínica: solo el grosor", 2)
parrafo("De 1.372 correlaciones, 33 sobrevivieron a la corrección. Su distribución por medida "
        "es el hallazgo:")

tabla(["Medida", "Pruebas", "Sobreviven FDR", "|ρ| máx"],
      [["Grosor", "260", "32", "0,702"],
       ["Área", "260", "1", "0,491"],
       ["Volumen", "592", "0", "0,551"],
       ["LGI", "260", "0", "0,421"]],
      resaltar_filas=(0,))
parrafo("Tabla 3. Con 592 pruebas el volumen no produce ninguna. El LGI tampoco, con 260. "
        "El grosor, con las mismas 260, produce 32.", size=9, color=TINTA2)

tabla(["Región", "Hemi", "Outcome", "ρ parcial", "p FDR"],
      [["Giro supramarginal", "der", "Niigata", "−0,70", "0,0010"],
       ["Giro supramarginal", "der", "DHI", "−0,69", "0,0012"],
       ["Prefrontal dorsolateral", "der", "DHI", "−0,68", "0,0012"],
       ["Postcentral", "izq", "DHI", "−0,68", "0,0012"],
       ["Temporal superior", "der", "DHI", "−0,65", "0,0018"],
       ["Occipital lateral", "der", "DHI", "−0,59", "0,0081"],
       ["Postcentral", "izq", "Niigata", "−0,59", "0,0145"],
       ["Temporal superior", "der", "Niigata", "−0,57", "0,0163"],
       ["Parietal inferior", "izq", "DHI", "−0,56", "0,0115"],
       ["Cingulada anterior", "izq", "DHI", "−0,55", "0,0142"]],
      resaltar_filas=(0,))
parrafo("Tabla 4. Correlación parcial de Spearman entre grosor cortical y severidad, dentro "
        "de pacientes (n ≈ 31), ajustada por edad, sexo y grupo. Menor grosor, mayor "
        "severidad. La red implicada excede las regiones de prioridad alta.",
        size=9, color=TINTA2)

figura(FIGS / "etapaB4" / "scatter_2_supramarginal_thickness_rh_Niigata.png", 4,
       "Grosor del giro supramarginal derecho y severidad sintomática.",
       "La recta se ajusta por grupo, nunca una sola global, para que se vea si la relación "
       "existe dentro de cada grupo o solo entre ellos. La correlación se replica dentro de "
       "cada grupo por separado —MPPP ρ = −0,70 (n = 14, p = 0,005) y vestibular ρ = −0,69 "
       "(n = 17, p = 0,002)—, de modo que no es un artefacto de mezclar grupos.",
       ancho_cm=11.5)

nota("CAUTELA OBLIGATORIA. Niigata y DHI correlacionan entre sí ρ = 0,72: miden esencialmente "
     "el mismo constructo y cuentan como un solo hallazgo, no como dos. En cambio la severidad "
     "clínica NO correlaciona con el desempeño en navegación (CSE–Niigata ρ = 0,26, p = 0,17), "
     "de modo que el eje clínico y el conductual sí son dimensiones independientes.")

figura(FIGS / "etapaD2_freesurfer" / "fig3_clusters_grosor_DHI.png", 5,
       "Clusters de asociación entre grosor cortical y severidad.",
       "Vertex-wise con la severidad como regresor continuo, dentro de pacientes. Azul = "
       "correlación negativa. Tres regiones coinciden exactamente con el análisis por región: "
       "temporal superior derecho (488 mm², CWP = 0,0002), postcentral izquierdo (336 mm², "
       "CWP = 0,0008) y prefrontal dorsolateral derecho.")

figura(FIGS / "etapaD2_freesurfer" / "fig4_doble_disociacion.png", 6,
       "La doble disociación, en una sola imagen.",
       "Arriba: los clusters de girificación que separan MPPP de pacientes vestibulares "
       "(vista dorsal, donde mejor se aprecia el cúmulo precentral izquierdo). Abajo: los "
       "clusters donde el grosor cortical escala con la severidad sintomática (vista "
       "lateral). Dos medidas de la misma corteza, en regiones parcialmente solapadas, "
       "capturando fenómenos distintos: rasgo y estado. Candidata a figura principal del "
       "manuscrito.")

titulo("3.4 Robustez", 2)
tabla(["Comprobación", "Resultado"],
      [["Réplica entre atlas", "r = 0,997 (DKT–DK), 0,947 (DKT–Destrieux), 0,954 (DK–Destrieux); "
                               "16/16 regiones conservan el signo"],
       ["Leave-one-out", "Ningún sujeto es decisivo: 0 de 36 reestimaciones pierden p < 0,05; "
                         "los ρ recorren rangos de 0,06–0,11"],
       ["Ajuste por ansiedad y depresión", "El tamaño de efecto no disminuye, aumenta "
                                           "(+0,07 a +0,12); lo que cae es el N"],
       ["Calidad de imagen", "SurfaceHoles no difiere entre grupos (p = 0,896)"]])
parrafo("Tabla 5. Las cuatro comprobaciones de robustez. La segunda responde directamente la "
        "objeción de que un coeficiente de 0,70 con n = 31 esté inflado por un caso "
        "influyente.", size=9, color=TINTA2)

titulo("3.5 Resultados negativos", 2)
parrafo("Se enumeran porque acotan lo que hay que explicar. La asimetría hemisférica no "
        "difiere entre grupos (0 de 68 índices en ambos diseños), lo que descarta la lectura "
        "lateralizada sugerida por la literatura previa. La covarianza estructural de la red "
        "no muestra diferencias que sobrevivan a la corrección. Las subestructuras (subcampos "
        "hipocampales, núcleos talámicos y amigdalinos) no producen ninguna familia "
        "enriquecida. El volumen cortical no produce ningún hallazgo en 484 correlaciones ni "
        "en el análisis de grupo.")

doc.add_page_break()

# %% ── DISCUSIÓN ────────────────────────────────────────────────────────────
titulo("4. Discusión", 1)

titulo("4.1 La doble disociación", 2)
tabla(["", "LGI", "Grosor cortical"],
      [["¿Diferencia MPPP de vestibular?", "Sí (d ≈ −0,9; 8 resultados; 7 clusters)",
        "No (0 de 136 dirigidas; 0,2× el azar)"],
       ["¿Correlaciona con conducta?", "No (0 de 260)", "Sí (ρ ≈ 0,45–0,60)"],
       ["¿Correlaciona con severidad?", "No", "Sí (ρ hasta −0,70; 32 supervivientes)"]])
parrafo("Tabla 6. Dos medidas de la misma corteza, en gran medida sobre las mismas regiones, "
        "con comportamientos ortogonales.", size=9, color=TINTA2)

parrafo("El resultado central no es ninguno de los dos hallazgos por separado, sino su "
        "relación. Es difícil de producir por azar: si todo fuera ruido, no esperaríamos que "
        "una medida concentrara toda la señal entre grupos y la otra toda la señal con "
        "severidad, apuntando ambas a las mismas regiones.")

titulo("4.2 Interpretación: rasgo y estado", 2)
parrafo("La girificación cortical se establece en el desarrollo temprano y es notablemente "
        "estable en la adultez; no es plausible que se modifique en los meses que dura un "
        "cuadro de MPPP. Un efecto en LGI apunta por tanto a un rasgo predisponente: una "
        "configuración cortical previa que hace a ciertos individuos más vulnerables a "
        "cronificar tras un evento vestibular agudo. Que el LGI no varíe con la severidad "
        "refuerza esta lectura.")
parrafo("El grosor cortical es plástico y responde a procesos adquiridos. Su correlación con "
        "la severidad actual, y su completa ausencia de diferencia entre grupos, lo sitúa como "
        "marcador de estado.")
parrafo("El modelo que sugieren los datos es de dos tiempos: una predisposición estructural "
        "determina quién cronifica; un proceso adquirido en la red vestibular cortical escala "
        "con cuán grave está.")

titulo("4.3 La limitación central", 2)
nota("EL CONTRASTE QUE SOBREVIVE ES CONTRA EL GRUPO VESTIBULAR, NO CONTRA LOS SANOS. El "
     "gradiente observado (MPPP −0,28 · Sano +0,04 · Vestibular +0,24) sitúa a los voluntarios "
     "sanos en posición intermedia. Caben dos lecturas incompatibles: (a) girificación "
     "reducida en MPPP como marcador de vulnerabilidad, o (b) girificación aumentada en el "
     "paciente vestibular que compensa bien, como marcador de reserva estructural. Con n = 10 "
     "sanos no es posible decidir entre ambas. Son dos artículos distintos, ambos publicables, "
     "con implicaciones clínicas opuestas. RECOMENDACIÓN: plantear ambas lecturas en la "
     "Discusión en lugar de elegir la más favorable; un revisor competente lo notará, y es "
     "mejor haberlo dicho primero.")

titulo("4.4 Otras limitaciones", 2)
parrafo("El tamaño muestral limita el diseño de tres grupos y ensancha los intervalos de "
        "confianza, varios de los cuales rozan el cero. El diseño es transversal: la "
        "interpretación rasgo/estado es una inferencia basada en las propiedades conocidas de "
        "cada medida, no una demostración longitudinal. Niigata y DHI no son independientes. "
        "Los coeficientes seleccionados de un barrido están sesgados al alza, aunque la "
        "replicación intra-grupo y el leave-one-out mitigan la objeción.")

titulo("4.5 Qué haría falta para cerrar la pregunta", 2)
parrafo("Primero, ampliar el brazo sano hasta n ≈ 25–30: es la única forma de resolver la "
        "ambigüedad direccional y convertiría un hallazgo ambiguo en uno direccional. Segundo, "
        "un seguimiento longitudinal de pacientes vestibulares agudos con imagen basal, para "
        "contrastar la hipótesis de rasgo predisponente de forma directa. Tercero, replicación "
        "independiente del eje grosor–severidad, que es el más robusto.")

doc.add_page_break()

# %% ── APÉNDICE ─────────────────────────────────────────────────────────────
titulo("Apéndice · Estado de las figuras", 1)
tabla(["Figura", "Estado", "Archivo o instrucción"],
      [["1 · Mapa anatómico de la red", "✓ FreeSurfer",
        "figs/etapaD2_freesurfer/fig1_mapa_red_DCNN.pdf"],
       ["2 · Forest de girificación", "✓ lista", "figs/sintesis/forest_LGI_todas_las_rois.pdf"],
       ["3 · Clusters de LGI", "✓ FreeSurfer",
        "figs/etapaD2_freesurfer/fig2_clusters_LGI.pdf"],
       ["4 · Dispersión grosor–Niigata", "✓ lista",
        "figs/etapaB4/scatter_2_supramarginal_thickness_rh_Niigata.pdf"],
       ["5 · Clusters de grosor", "✓ FreeSurfer",
        "figs/etapaD2_freesurfer/fig3_clusters_grosor_DHI.pdf"],
       ["6 · Panel de la disociación", "✓ FreeSurfer",
        "figs/etapaD2_freesurfer/fig4_doble_disociacion.pdf"]])

parrafo("Las seis figuras están generadas. Las de superficie se renderizaron con freeview en "
        "modo batch (notebooks/etapaD2_figuras_freesurfer.py); las versiones previas hechas "
        "con nilearn se conservan en figs/etapaD/ y figs/etapaB5/ como respaldo. Todas "
        "existen en PNG a 200 dpi y PDF vectorial.", size=9.5, color=TINTA2)

parrafo("Documento generado automáticamente desde notebooks/generar_word_paper.py. "
        "Regenerable; las ediciones manuales sobre el .docx se pierden al regenerar. "
        "La fuente de verdad es docs/PAPER_BORRADOR.md.", cursiva=True, size=9,
        color=TINTA2)

# %% ── guardar ──────────────────────────────────────────────────────────────
doc.save(SALIDA)
print(f"→ {SALIDA}  ({SALIDA.stat().st_size/1024:.0f} KB)")

# copia junto a la documentación del proyecto
destino = cfg.ONEDRIVE / "PAPER_BOSQUEJO.docx"
try:
    shutil.copy2(SALIDA, destino)
    print(f"→ copia en {destino}")
except Exception as e:
    print(f"⚠️ no se pudo copiar a OneDrive: {e}")
